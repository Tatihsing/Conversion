"""
build_docx.py
根據 MEETING dict 產出 Word (.docx) 會議記錄（詳細內文），方便後續編輯修改
"""

import sys
import re
import subprocess
from pathlib import Path
from datetime import datetime


def _ensure_docx():
    """自動安裝 python-docx"""
    try:
        import docx
        return True
    except ImportError:
        print("[INSTALL] 安裝 python-docx 中...")
        result = subprocess.run(
            [sys.executable, "-m", "pip", "install", "python-docx"],
            capture_output=True, text=True, encoding='utf-8', errors='replace'
        )
        if result.returncode != 0:
            print("[WARN] python-docx 安裝失敗，跳過 Word 檔產出")
            return False
        print("[OK] python-docx 安裝完成")
        return True


def build_docx(summary, meeting, output_path, gen_time=None):
    """
    產出 Word 文件（下半部詳細內文 + Action Items）
    """
    if not _ensure_docx():
        return None

    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    if gen_time is None:
        gen_time = datetime.now().strftime('%Y-%m-%d %H:%M')

    doc = Document()

    # ── 設定預設字型 ──
    style = doc.styles['Normal']
    font = style.font
    font.name = '微軟正黑體'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # 設定段落間距
    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(6)
    paragraph_format.line_spacing = 1.5

    # ── 大標題 ──
    title = doc.add_heading(level=0)
    title_run = title.add_run(summary.get("big_title", "會議記錄"))
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
    title_run.font.name = '微軟正黑體'

    # ── 副標題 + 目標 ──
    sub = summary.get("sub_title", "")
    if sub:
        p = doc.add_paragraph()
        run = p.add_run(sub)
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.name = '微軟正黑體'

    obj = summary.get("objective", "")
    if obj:
        p = doc.add_paragraph()
        run = p.add_run(obj)
        run.font.size = Pt(11)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        run.font.name = '微軟正黑體'

    doc.add_paragraph()  # 空行

    # ── 詳細段落 ──
    sections = meeting.get("sections", [])
    for sec in sections:
        heading_text = sec.get("heading", "")

        # 段落標題
        h = doc.add_heading(level=1)
        h_run = h.add_run(heading_text)
        h_run.font.size = Pt(16)
        h_run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
        h_run.font.name = '微軟正黑體'

        # 段落內容
        for para_text in sec.get("paras", []):
            p = doc.add_paragraph()
            run = p.add_run(para_text)
            run.font.size = Pt(12)
            run.font.name = '微軟正黑體'
            p.paragraph_format.first_line_indent = Cm(0.8)

        # 條列項目
        bullets = sec.get("bullets", [])
        for level, text in bullets:
            if isinstance(level, str):
                try:
                    level = int(level)
                except ValueError:
                    level = 1

            if level <= 1:
                p = doc.add_paragraph(style='List Bullet')
            else:
                p = doc.add_paragraph(style='List Bullet 2')
            run = p.add_run(str(text))
            run.font.size = Pt(12)
            run.font.name = '微軟正黑體'

        # 段落結語
        closing = sec.get("closing", "")
        if closing:
            p = doc.add_paragraph()
            run = p.add_run(closing)
            run.font.size = Pt(11)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            run.font.name = '微軟正黑體'
            # 加底色效果（用段落邊框模擬）
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(8)

    # ── Action Items ──
    action_items = meeting.get("action_items", [])
    if action_items:
        h = doc.add_heading(level=1)
        h_run = h.add_run("Action Items 詳細清單")
        h_run.font.size = Pt(16)
        h_run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
        h_run.font.name = '微軟正黑體'

        p = doc.add_paragraph()
        run = p.add_run("以下行動項依責任人彙整，所有未具體載明期限者均以 [TBD] 標註。")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        run.font.name = '微軟正黑體'

        for person in action_items:
            owner = person.get("owner", "")
            # 負責人名稱
            p = doc.add_paragraph()
            run = p.add_run(f"■ {owner}")
            run.font.size = Pt(13)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
            run.font.name = '微軟正黑體'

            for idx, item in enumerate(person.get("items", []), 1):
                clean_item = re.sub(r'^\[.?\]\s*', '', item)
                p = doc.add_paragraph(style='List Number')
                run = p.add_run(clean_item)
                run.font.size = Pt(12)
                run.font.name = '微軟正黑體'

    # ── 頁尾 ──
    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run(f"本文件由 AI 自動產生 · 生成時間：{gen_time} · 內容依逐字稿整理，如有出入以錄音為準")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    run.font.name = '微軟正黑體'

    # ── 儲存 ──
    doc.save(str(output_path))
    print(f"[OK] Word 已產出：{output_path}")
    return str(output_path)
